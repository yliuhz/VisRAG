
import pandas as pd
import polars as pl
import numpy as np
import os

from PIL import Image
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches

import matplotlib.pyplot as plt
import seaborn as sns


def write_to_docx(df: pd.DataFrame, filename: str, topk: int):

    df = df.sort_values(by=['query-id', 'rank'], ascending=True)

    print(list(df.columns))

    document = Document()
    print_n = 100

    for i in range(0,print_n,topk):
        paragraph = document.add_paragraph()

        query_ids = None
        for idx in range(i, i+topk): ## image idx
            image = BytesIO(df['image'].iloc[idx]["bytes"])

            run = paragraph.add_run()
            run.add_picture(image, width=Inches(2.5))

        for idx in range(i, i+topk): 
            rank, score = df['rank'].iloc[idx], df['score'].iloc[idx]

            gt_doc_ids = []
            for col in df.columns:
                if 'corpus' in col:
                    gt_doc_ids.append(df[col].iloc[idx])

            doc_id = df['doc-id'].iloc[idx]
            document.add_paragraph(
                f"{rank}\t{score}\t{doc_id in gt_doc_ids}", style='List Bullet'
            )

            if query_ids is None:
                query_ids = df['query-id'].iloc[idx]
            else:
                assert query_ids == df['query-id'].iloc[idx], f"{query_ids}\t{df['query-id'].iloc[idx]}"

        
                
        query = df["query"].iloc[i]
        document.add_paragraph(
            query, style='List Bullet'
        )

        options = df["options"].iloc[i]
        if options is not None:
            for option in options:
                document.add_paragraph(
                    option, style='List Bullet'
                )


        answer = df["answer"].iloc[i]
        document.add_paragraph(
            f"Answer: {answer}", style='List Bullet'
        )

        document.add_page_break()
    document.save(filename)

def load_parquet(root=None, return_type='pd'):
    """
    Load data from Parquet files with optional filtering on date_id, time_id, and selected columns.

    Parameters:
    - date_id_range (tuple, optional): Range of date_id to filter (start, end). Default is None, which means all dates.
    - time_id_range (tuple, optional): Range of time_id to filter (start, end). Default is None, which means all times.
    - columns (list, optional): List of columns to load. Default is None, which means all columns.
    - return_type (str, optional): Type of data to return ('pl' for Polars DataFrame or 'pd' for Pandas DataFrame). Default is 'pl'.

    Returns:a
    - pl.DataFrame or pd.DataFrame: The filtered data as a Polars or Pandas DataFrame.
    """
    # Load data using Polars lazy loading (scan_parquet)
    data = pl.scan_parquet(f"{root}")

    # Collect the data to execute the lazy operations
    if return_type == 'pd':
        return data.collect().to_pandas()
    else:
        return data.collect()

if __name__ == "__main__":

    root = f"/ssddata/liuyue/github/VisRAG/data/checkpoints/eval-2024-11-06-174337-maxq-512-maxp-2048-bsz-16-pooling-wmean-attention-causal-gpus-per-node-8"
    file_temp = "test.{}.trec"
    datasets = [
        "ArxivQA",
        "ChartQA",
        "MP-DocVQA",
        "InfoVQA",
        "PlotQA",
        "SlideVQA",
    ]
    GPUS=8
    topk = 3

    for dataset in datasets:
        print(f'Dataset: {dataset}')

        df = pd.DataFrame()
        for i in range(GPUS):
            path = f"{root}/{dataset}/{file_temp.format(i)}"

            df_s = pd.read_csv(path, sep='\t', header=None, names=['query-id', 'Q0', 'doc-id', 'rank', 'score', 'run-id'])
            df = pd.concat([df,df_s], axis=0)

            print(f'Read {len(df_s)} lines, {df_s.shape}, {df.shape}')


        df.columns = ['query-id', 'Q0', 'doc-id', 'rank', 'score', 'run-id']
        df['rank'] = pd.to_numeric(df['rank'])
        df = df[df['rank'] <= 3]

        #### load data 
        root2 = "/ssddata/liuyue/github/VisRAG/qa_datasets"
        file_temp2 = "VisRAG-Ret-Test-{}"
        outdir = "/ssddata/liuyue/github/VisRAG/data/checkpoints/eval-2024-11-06-174337-maxq-512-maxp-2048-bsz-16-pooling-wmean-attention-causal-gpus-per-node-8/{}"  

    # for data in datasets:
        corpus_df = load_parquet(f"{root2}/{file_temp2.format(dataset)}/corpus/*.parquet")
        qrel_df = load_parquet(f"{root2}/{file_temp2.format(dataset)}/qrels/*.parquet")
        query_df = load_parquet(f"{root2}/{file_temp2.format(dataset)}/queries/*.parquet")

        if dataset == "SlideVQA": ### there are duplicates in query-id
            ### q1  c1                  q1  c1  c2
            ### q1  c2          -->     q2  c1  NaN
            ### q2  c1
            #### https://stackoverflow.com/questions/66785979/remove-duplicates-and-add-some-column-using-python-pandas#:~:text=(df.drop_duplicates()%0A%20%20%20.assign(col%3Dlambda%20x%3A%20x.groupby(%22Name%22).cumcount())%0A%20%20%20.pivot(index%3D%27Name%27%2C%20columns%3D%27col%27%2C%20values%3D%27Email%27)%0A%20%20%20.add_prefix(%27Email_%27).reset_index()%0A)
            qrel_df = (qrel_df.assign(col=lambda x: x.groupby("query-id").cumcount())
            .pivot(index='query-id', columns='col', values='corpus-id')
            .add_prefix('corpus_').reset_index()
            )

        ## https://pandas.pydata.org/docs/reference/api/pandas.DataFrame.join.html#:~:text=Another%20option%20to%20join%20using%20the%20key%20columns%20is%20to%20use%20the%20on%20parameter.%20DataFrame.join%20always%20uses%20other%E2%80%99s%20index%20but%20we%20can%20use%20any%20column%20in%20df.
        data_df = df.join(corpus_df.set_index('corpus-id'), on='doc-id', how='left')
        data_df = data_df.join(query_df.set_index('query-id'), on='query-id', how='left')
        data_df = data_df.join(qrel_df.set_index('query-id'), on='query-id', how='left', rsuffix="_right")

        print(data_df.shape, query_df.shape, qrel_df.shape)

        # write_to_docx(data_df, f"{outdir.format(dataset)}/{dataset}.docx", topk)

        # df = data_df[['query-id', 'doc-id', 'rank', 'score']]
        df = data_df
        df2 = df[df['rank']==1].join(df[df['rank']==3].set_index('query-id'), on='query-id', how='inner', lsuffix='_left', rsuffix='_right')

        df2['bias'] = df2['score_left'] - df2['score_right'] ###
        df2['r_bias'] = df2['bias'] / df2['score_left'] ###
        
        df2['correct'] = np.zeros(len(df2))
        df2 = df2.fillna(0)
        print(df2.head(2))
        for col in df2.columns:
            if 'corpus' in col:
                df2['correct'] += (df2['doc-id_left'] == df2[col]) ## doc-id_left: rank1 retrieval is correct
        df2['correct'] = (df2['correct'] > 0) ###

        df3 = df2.groupby(by=['query-id']).sum()
        df3['correct'] = (df3['correct'] > 0) ###

        fig = plt.figure(figsize=(10,10), dpi=300)
        ax = plt.gca()
        sns.scatterplot(
            data = df3,
            x = 'score_left',
            y = 'r_bias',
            hue = 'correct',
            ax=ax,
        )
        ax.set_ylim(bottom=-0.05)
        ax.set_xlabel(f'Rank1 score')
        ax.set_ylabel(f'(Rank1 score - Rank3 score) / Rank1 score')
        ax.set_title(f'{dataset}')
        fig.savefig(f"{outdir.format(dataset)}/{dataset}.png")