import pandas as pd
import polars as pl
import numpy as np
import os

from PIL import Image
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches

def write_to_docx(df: pd.DataFrame, filename: str, true: int):

    df = df[df["correct"] == true].iloc[:100, :]
    df = df.sample(frac=1).reset_index(drop=True)
    print(list(df.columns))

    document = Document()
    
    for i in range(len(df)):
        paragraph = document.add_paragraph()
        for col in df.columns:
            if "image" in col:

                try:
                    image = BytesIO(df[col].iloc[i]["bytes"])

                    run = paragraph.add_run()
                    run.add_picture(image, width=Inches(2.5))
                except: ### different queries have different number of references
                    continue

        query = df["query"].iloc[i]
        document.add_paragraph(
            query, style='List Bullet'
        )

        options = df["options"].iloc[i]
        if options is not None:
            for option in options:
                document.add_paragraph(
                    option, style='List Bullet'
                )


        answer = df["answer"].iloc[i]
        document.add_paragraph(
            f"Answer: {answer}", style='List Bullet'
        )

        resp = df["respond"].iloc[i]
        document.add_paragraph(
            f"Respond: {resp}", style='List Bullet'
        )

        correct = df["correct"].iloc[i]
        document.add_paragraph(
            f'Correct: {bool(correct)}', style='List Bullet'
        )

        document.add_page_break()
    document.save(filename)


def load_parquet(root=None, return_type='pd'):
    """
    Load data from Parquet files with optional filtering on date_id, time_id, and selected columns.

    Parameters:
    - date_id_range (tuple, optional): Range of date_id to filter (start, end). Default is None, which means all dates.
    - time_id_range (tuple, optional): Range of time_id to filter (start, end). Default is None, which means all times.
    - columns (list, optional): List of columns to load. Default is None, which means all columns.
    - return_type (str, optional): Type of data to return ('pl' for Polars DataFrame or 'pd' for Pandas DataFrame). Default is 'pl'.

    Returns:a
    - pl.DataFrame or pd.DataFrame: The filtered data as a Polars or Pandas DataFrame.
    """
    # Load data using Polars lazy loading (scan_parquet)
    data = pl.scan_parquet(f"{root}")

    # Collect the data to execute the lazy operations
    if return_type == 'pd':
        return data.collect().to_pandas()
    else:
        return data.collect()
    
if __name__ == "__main__":

    ### extract answer correct or not from the log
    root = "/ssddata/liuyue/github/VisRAG/logs/generator/MiniCPMV2.6"
    file_temp = "eval_g_{}_1.log"
    datasets = [
        "ArxivQA",
        "ChartQA",
        "MP-DocVQA",
        "InfoVQA",
        "PlotQA",
        "SlideVQA",
    ]

    data2correctlist = {k:[] for k in datasets}
    for data in datasets:
        path = f"{root}/{file_temp.format(data)}"

        correct = []
        queries = []
        responds = []
        with open(f"{path}", "r") as f:
            lines = f.readlines()
            for line in lines:
                if ", True" in line:
                    correct.append(1)
                elif ", False" in line:
                    correct.append(0)
        data2correctlist[data] = {"correct": correct}

    #### extract query_id from history jsonl, it has the same query order with that of log
    root = "/ssddata/liuyue/github/VisRAG/data/checkpoints/generator/MiniCPMV2.6/upper_bound"
    file_temp = "MiniCPMV2.6_upper_bound_{}_oracle_history.jsonl"
    for data in datasets:
        path = f"{root}/{file_temp.format(data)}"

        qids, queries, resps = [], [], []
        jsonlines = pd.read_json(path_or_buf=path, lines=True)
        # for line in jsonlines:
        #     qids.append(line['qid'])
        #     queries.append(line['query'])
        #     resps.append(line['preprocessed_responds'])
        qids = jsonlines['qid'].tolist()
        queries = jsonlines['query'].tolist()
        resps = jsonlines['preprocessed_responds'].tolist()

        ## merge two dicts, https://stackoverflow.com/questions/38987/how-do-i-merge-two-dictionaries-in-a-single-expression-in-python#:~:text=discussed%20here)%3A-,z%20%3D%20x%20%7C%20y,-In%20Python%203.5
        data2correctlist[data] = data2correctlist[data] | {'query-id': qids, 'respond': resps}


    
    #### load data 
    root = "/ssddata/liuyue/github/VisRAG/QAdataset"
    file_temp = "VisRAG-Ret-Test-{}"
    outdir = "/ssddata/liuyue/github/VisRAG/data/checkpoints/generator/MiniCPMV2.6/upper_bound"
    for data in datasets:
        print(f"{data}")

        corpus_df = load_parquet(f"{root}/{file_temp.format(data)}/corpus/*.parquet")
        qrel_df = load_parquet(f"{root}/{file_temp.format(data)}/qrels/*.parquet")
        query_df = load_parquet(f"{root}/{file_temp.format(data)}/queries/*.parquet")
        
        if data == "SlideVQA": ### there are duplicates in query-id
            ### q1  c1                  q1  c1  c2
            ### q1  c2          -->     q2  c1  NaN
            ### q2  c1
            #### https://stackoverflow.com/questions/66785979/remove-duplicates-and-add-some-column-using-python-pandas#:~:text=(df.drop_duplicates()%0A%20%20%20.assign(col%3Dlambda%20x%3A%20x.groupby(%22Name%22).cumcount())%0A%20%20%20.pivot(index%3D%27Name%27%2C%20columns%3D%27col%27%2C%20values%3D%27Email%27)%0A%20%20%20.add_prefix(%27Email_%27).reset_index()%0A)
            qrel_df = (qrel_df.assign(col=lambda x: x.groupby("query-id").cumcount())
            .pivot(index='query-id', columns='col', values='corpus-id')
            .add_prefix('corpus_').reset_index()
            )


        data_df = query_df.join(qrel_df.set_index("query-id"), on="query-id", how="inner")

        if data == "SlideVQA":  
            for col in qrel_df.columns:
                if 'corpus' in col:
                    data_df = data_df.join(corpus_df.set_index("corpus-id"), on=col, how="left", lsuffix="_left", rsuffix="_right")

        else:
            data_df = data_df.join(corpus_df.set_index("corpus-id"), on="corpus-id", how="inner")

        correct_df = pd.DataFrame(data2correctlist[data])

        data_df2 = data_df.join(correct_df.set_index("query-id"), on="query-id", how="inner")
        # data_df = data_df["corpus-id", "image", "query-id", "query", "answer", "is_numerical", "correct"]

        print(data_df2.shape, data_df.shape, query_df.shape, correct_df.shape, qrel_df.shape)
        # data_df2.to_csv(f"{outdir}/{data}.csv")

        write_to_docx(data_df2, f"{outdir}/{data}.docx", 0)
        write_to_docx(data_df2, f"{outdir}/{data}-correct.docx", 1)
        


        # print(corpus_df.shape, corpus_df.drop_duplicates(['corpus-id']).shape)
        # print(qrel_df.shape, qrel_df.drop_duplicates(['query-id']).shape, qrel_df.drop_duplicates(['corpus-id']).shape)
        # print(query_df.shape, query_df.drop_duplicates(['query-id']).shape, query_df.drop_duplicates(['query']).shape)
        # print(correct_df.shape, correct_df.drop_duplicates(['query']).shape)
        print(f"="*30)


    pass